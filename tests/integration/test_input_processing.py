from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from harness.adapters.memory import (
    InMemoryArtifactStore,
    InMemoryInputArtifactRepository,
    InMemoryThreadFileRepository,
)
from harness.application.file_catalog import FileCatalogService
from harness.application.input_artifacts import InputArtifactService
from harness.core.models import ExecutionIdentity, ThreadFileKind
from harness.inputs.processors import DefaultInputProcessor

NOW = datetime(2026, 7, 13, tzinfo=UTC)


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Private Brief", level=1)
    document.add_paragraph("The launch code is blue-42.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_original_and_processed_files_are_staged_and_cataloged(
    tmp_path: Path,
) -> None:
    sequence = 0

    def ids(prefix: str) -> str:
        nonlocal sequence
        sequence += 1
        return f"{prefix}_{sequence}"

    catalog = FileCatalogService(
        InMemoryThreadFileRepository(), clock=lambda: NOW, id_generator=ids
    )
    service = InputArtifactService(
        repository=InMemoryInputArtifactRepository(),
        store=InMemoryArtifactStore(),
        id_generator=ids,
        clock=lambda: NOW,
        processor=DefaultInputProcessor(),
        file_catalog=catalog,
    )
    original_bytes = docx_bytes()
    uploaded = await service.upload(
        tenant_id="tenant-a",
        user_id="alice",
        name="brief.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        content=original_bytes,
    )
    identity = ExecutionIdentity(
        tenant_id="tenant-a",
        user_id="alice",
        project_id="agent-a",
        session_id="session-a",
        run_id="run-a",
        agent_name="agent-a",
        agent_version="1.0.0",
    )

    staged = await service.stage_for_run(
        tenant_id="tenant-a",
        user_id="alice",
        input_artifact_ids=[uploaded.input_artifact_id],
        workspace=tmp_path,
        identity=identity,
    )

    assert staged[0].path.startswith("inputs/original/")
    assert staged[0].processed_paths[0].startswith("inputs/processed/")
    markdown = (tmp_path / staged[0].processed_paths[0]).read_text()
    assert "# Private Brief" in markdown
    assert "blue-42" in markdown
    assert (tmp_path / staged[0].path).read_bytes() == original_bytes
    files = await catalog.list_for_thread(identity)
    assert [file.kind for file in files] == [
        ThreadFileKind.ORIGINAL,
        ThreadFileKind.DERIVED,
    ]
    assert files[1].parent_file_id == files[0].file_id


@pytest.mark.asyncio
async def test_processor_failure_preserves_original_and_records_failure(
    tmp_path: Path,
) -> None:
    class FailingProcessor:
        def process(self, *, name: str, media_type: str, content: bytes) -> None:
            del name, media_type, content
            raise ValueError("private parser detail")

    sequence = 0

    def ids(prefix: str) -> str:
        nonlocal sequence
        sequence += 1
        return f"{prefix}_{sequence}"

    catalog = FileCatalogService(
        InMemoryThreadFileRepository(), clock=lambda: NOW, id_generator=ids
    )
    service = InputArtifactService(
        repository=InMemoryInputArtifactRepository(),
        store=InMemoryArtifactStore(),
        id_generator=ids,
        clock=lambda: NOW,
        processor=FailingProcessor(),  # type: ignore[arg-type]
        file_catalog=catalog,
    )
    uploaded = await service.upload(
        tenant_id="tenant-a",
        user_id="alice",
        name="broken.bin",
        media_type="application/octet-stream",
        content=b"immutable-original",
    )
    identity = ExecutionIdentity(
        tenant_id="tenant-a",
        user_id="alice",
        project_id="agent-a",
        session_id="session-a",
        run_id="run-a",
        agent_name="agent-a",
        agent_version="1.0.0",
    )

    staged = await service.stage_for_run(
        tenant_id="tenant-a",
        user_id="alice",
        input_artifact_ids=[uploaded.input_artifact_id],
        workspace=tmp_path,
        identity=identity,
    )

    assert (tmp_path / staged[0].path).read_bytes() == b"immutable-original"
    assert staged[0].processed_paths == ()
    files = await catalog.list_for_thread(identity)
    assert files[0].metadata["processing_status"] == "failed"
    assert files[0].metadata["processing_error_code"] == "ValueError"
    assert "private parser detail" not in repr(files)


@pytest.mark.asyncio
async def test_restaging_replaces_read_only_input_from_restored_workspace(
    tmp_path: Path,
) -> None:
    sequence = 0

    def ids(prefix: str) -> str:
        nonlocal sequence
        sequence += 1
        return f"{prefix}_{sequence}"

    service = InputArtifactService(
        repository=InMemoryInputArtifactRepository(),
        store=InMemoryArtifactStore(),
        id_generator=ids,
        clock=lambda: NOW,
    )
    uploaded = await service.upload(
        tenant_id="tenant-a",
        user_id="alice",
        name="image.png",
        media_type="image/png",
        content=b"image-content",
    )

    first = await service.stage_for_run(
        tenant_id="tenant-a",
        user_id="alice",
        input_artifact_ids=[uploaded.input_artifact_id],
        workspace=tmp_path,
    )
    target = tmp_path / first[0].path
    assert target.stat().st_mode & 0o777 == 0o444

    second = await service.stage_for_run(
        tenant_id="tenant-a",
        user_id="alice",
        input_artifact_ids=[uploaded.input_artifact_id],
        workspace=tmp_path,
    )

    assert second[0].path == first[0].path
    assert target.read_bytes() == b"image-content"
    assert target.stat().st_mode & 0o777 == 0o444
